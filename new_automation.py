"""
Y-Axis LP Reporting Automation  — v4
=======================================
Excel CAS  →  Word template (per LP)  →  PDF (grouped by Entity)

Requirements (install once):
    pip install openpyxl python-docx pywin32

Run:
    python lp_report_automation.py

CHANGES FROM v3:
  • Outlook draft step removed entirely — pipeline now ends at PDF creation.
  • PDFs are grouped into a subfolder per Entity name under OUTPUT_FOLDER,
    e.g. Output_PDFs/Ram Group Partners, LP/001_ABC_Holdings_LLC_Q1 2026.pdf
"""

import os, re, sys, shutil, openpyxl
from docx import Document
from datetime import datetime

# ─── CONFIGURATION ────────────────────────────────────────────────────────────
CAS_FILE      = r"CAS_Data\test_NAV.xlsx"
CAS_SHEET     = "Summary"
WORD_TEMPLATE = r"LP_Report_Template_v3.docx"
OUTPUT_FOLDER = r"Output_PDFs"
QUARTER_LABEL = "Q1 2026"

# ─── COLUMN → TEMPLATE TOKEN MAP ──────────────────────────────────────────────
COLUMN_MAP = {
    "{{INVESTOR_NAME}}":         "Investor",
    "{{ENTITY}}":                "Entity",
    "{{INVESTOR_EMAIL}}":        "Investor Email",
    "{{INVESTOR_CC_EMAIL}}":     "Investor CC Email",
    "{{DATE}}":                  "Date",
    "{{INVESTED_CAPITAL}}":      "Invested Capital",
    "{{ITD_BEGINNING_NAV}}":     "ITD Beginning NAV",
    "{{ITD_CONTRIBUTIONS}}":     "ITD Contributions",
    "{{ITD_CHANGE_UNREALIZED}}": "ITD Change in Unrealized App Dep",
    "{{ITD_RETURN_OF_CAPITAL}}": "ITD Return of Capital",
    "{{ITD_RETURN_ON_CAPITAL}}": "ITD Return on Capital",
    "{{ITD_ENDING_NAV}}":        "ITD Ending NAV",
    "{{YTD_BEGINNING_NAV}}":     "YTD beginning NAV",
    "{{YTD_CONTRIBUTIONS}}":     "YTD Contributions",
    "{{YTD_CHANGE_UNREALIZED}}": "YTD Change in Unrealized App Dep",
    "{{YTD_RETURN_OF_CAPITAL}}": "YTD Return of Capital",
    "{{YTD_RETURN_ON_CAPITAL}}": "YTD Return on Capital",
    "{{YTD_ENDING_NAV}}":        "YTD Ending NAV",
    "{{QTD_BEGINNING_NAV}}":     "QTD Beginning NAV",
    "{{QTD_CONTRIBUTIONS}}":     "QTD Contributions",
    "{{QTD_CHANGE_UNREALIZED}}": "QTD Change in Unrealized App Dep",
    "{{QTD_RETURN_OF_CAPITAL}}": "QTD Return of Capital",
    "{{QTD_RETURN_ON_CAPITAL}}": "QTD Return on Capital ",  # trailing space in Excel
    "{{QTD_ENDING_NAV}}":        "QTD Ending NAV",
}

# ─── HELPERS ──────────────────────────────────────────────────────────────────

# LINE 109 — FIX Problem 3a, 3b & 3d:
#   • Rounds all values to integer (no .00)
#   • Zero / blank / dash / None  →  "-"  (no "$0" shown)
#   • Negative numbers            →  "($2,313)"  not  "-$2,313"
def fmt(v):
    """
    Format a NAV/financial value for display:
      - None / "" / "-" / 0 / 0.0  →  "-"
      - Positive number             →  "$12,345"
      - Negative number             →  "($2,313)"   ← accounting parentheses format
      - Non-numeric string          →  returned as-is
    """
    if v is None or str(v).strip() in ("", "-", "—"):
        return "-"
    try:
        num = float(v)
        # LINE 120 — zero activity shows as "-" not "$0"
        if num == 0:
            return "-"
        # LINE 122 — FIX Problem 3d: negative → accounting parentheses format
        if num < 0:
            return f"(${abs(int(round(num))):,})"
        return f"${int(round(num)):,}"
    except (ValueError, TypeError):
        return str(v)


# LINE 126 — FIX Problem 3b & 3d: Invested Capital rounded integer, negatives in parens
def fmt_capital(v):
    """Invested Capital: rounded integer, zero → '-', negative → '($n)'."""
    if v is None or str(v).strip() in ("", "-", "—"):
        return "-"
    try:
        num = float(v)
        if num == 0:
            return "-"
        if num < 0:
            return f"(${abs(int(round(num))):,})"
        return f"${int(round(num)):,}"
    except (ValueError, TypeError):
        return str(v) if v else "-"


# LINE 130 — FIX Problem 3c: date only, strip any time component
def fmt_date(v):
    """
    Return only the date portion of the CAS date field.
    Handles: datetime objects, 'March 31,2026', '2026-03-31', '2026-03-31 00:00:00'
    """
    if v is None or str(v).strip() == "":
        return ""
    # Already a datetime/date object
    if hasattr(v, "strftime"):
        return v.strftime("%B %d, %Y")          # e.g. "March 31, 2026"
    s = str(v).strip()
    # Try parsing common formats and normalise
    for fmt_str in ("%B %d,%Y", "%B %d, %Y", "%Y-%m-%d %H:%M:%S",
                    "%Y-%m-%d", "%m/%d/%Y", "%d-%b-%Y"):
        try:
            return datetime.strptime(s, fmt_str).strftime("%B %d, %Y")
        except ValueError:
            continue
    # Fallback: strip anything after a space that looks like a time (HH:MM)
    s = re.sub(r"\s+\d{1,2}:\d{2}(:\d{2})?(\.\d+)?$", "", s).strip()
    return s


def fmt_date_numeric(v):
    """
    Same parsing as fmt_date(), but returns MM.DD.YYYY for use in filenames
    (fmt_date's "March 31, 2026" form is for the document body/template).
    """
    if v is None or str(v).strip() == "":
        return ""
    if hasattr(v, "strftime"):
        return v.strftime("%m.%d.%Y")
    s = str(v).strip()
    for fmt_str in ("%B %d,%Y", "%B %d, %Y", "%Y-%m-%d %H:%M:%S",
                    "%Y-%m-%d", "%m/%d/%Y", "%d-%b-%Y"):
        try:
            return datetime.strptime(s, fmt_str).strftime("%m.%d.%Y")
        except ValueError:
            continue
    s = re.sub(r"\s+\d{1,2}:\d{2}(:\d{2})?(\.\d+)?$", "", s).strip()
    return s


def read_cas(path, sheet):
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb[sheet]
    rows = list(ws.iter_rows(values_only=True))
    headers = [str(h).strip() if h else "" for h in rows[0]]
    return [{headers[i]: rows[r][i] for i in range(len(headers))}
            for r in range(1, len(rows)) if any(rows[r])]


# LINE 170 — Column A "Number" is unique by nature, so it alone is used as
# the internal key for temp-file naming / dict lookups (no longer combined
# with the investor name).
def unique_key(lp_number):
    """
    Zero-padded string form of the Number column, e.g. 1 -> "001".
    Used only for internal uniqueness (temp .docx name, pdf_map key) —
    NOT part of the final PDF filename, which is Entity/Date/Investor based.
    """
    return str(int(lp_number)).zfill(3) if lp_number not in (None, "") else "000"


def safe_folder_name(entity_name):
    """
    Sanitize an Entity name into a filesystem-safe folder name.
    Only strips characters Windows actually forbids in folder names
    (commas/periods are fine in a folder name, unlike in a filename component).
    """
    s = re.sub(r'[\\/*?:"<>|]', "_", str(entity_name)).strip()
    return s if s else "Unknown_Entity"


def _clean_path_component(text, fallback):
    """Strip only the characters Windows actually forbids; keep commas/periods."""
    s = re.sub(r'[\\/*?:"<>|]', "_", str(text)).strip() if text else ""
    return s if s else fallback


# Legal-entity suffixes to drop from the FILENAME only (folder name keeps
# the full entity name). Only strips when the last word is actually one of
# these — never blindly chops the last word, since "Delhi Group Partners"
# has no such suffix and must stay intact.
ENTITY_SUFFIXES_TO_STRIP = {
    "LP", "L.P.", "LLC", "L.L.C.", "LLP", "L.L.P.", "INC", "INC.", "LTD", "LTD.",
}


def strip_entity_suffix(text):
    """
    Remove a trailing legal-entity suffix (LP, LLC, Inc, Ltd, ...) and any
    leftover trailing comma. Leaves the name unchanged if the last word
    isn't a recognised suffix.
    e.g. "Ram Group Partners, LP" -> "Ram Group Partners"
         "Delhi Group Partners"   -> "Delhi Group Partners"  (unchanged)
    """
    s = str(text).strip()
    parts = s.rsplit(None, 1)  # split off the last whitespace-separated word
    if len(parts) == 2 and parts[1].strip(",.").upper() in ENTITY_SUFFIXES_TO_STRIP:
        s = parts[0].rstrip(",").strip()
    return s


def build_pdf_filename(entity, date_val, investor_name):
    """
    "[Full Entity Name, suffix stripped] CAS - MM.DD.YYYY - [Investor Name].pdf"
    e.g. "Ram Group Partners CAS - 03.31.2026 - ABC Holdings, LLC.pdf"
    """
    entity_clean   = _clean_path_component(entity, "Unknown Entity")
    entity_clean   = strip_entity_suffix(entity_clean)
    investor_clean = _clean_path_component(investor_name, "Unknown Investor")
    date_str       = fmt_date_numeric(date_val)
    entity_label   = f"{entity_clean} CAS" if entity_clean else "CAS"
    parts = [p for p in (entity_label, date_str, investor_clean) if p]
    return " - ".join(parts) + ".pdf"


def replace_in_para(para, tokens):
    """Word may split a merge field across runs — rebuild, replace, rewrite."""
    full = "".join(r.text for r in para.runs)
    if any(tok in full for tok in tokens):
        for tok, val in tokens.items():
            full = full.replace(tok, val)
        if para.runs:
            para.runs[0].text = full
            for r in para.runs[1:]:
                r.text = ""


def fill_doc(template_path, tokens):
    doc = Document(template_path)
    for para in doc.paragraphs:
        replace_in_para(para, tokens)
    for tbl in doc.tables:
        for row in tbl.rows:
            for c in row.cells:
                for para in c.paragraphs:
                    replace_in_para(para, tokens)
    for sec in doc.sections:
        for hdr_ftr in (sec.header, sec.footer):
            for para in hdr_ftr.paragraphs:
                replace_in_para(para, tokens)
    return doc


# LINE 218 — FIX Problem 1 (v2): Word COM timing/release race eliminated.
#   Root cause: win32com.client.Dispatch() can silently attach to a Word
#   instance still mid-shutdown from the PREVIOUS LP (Word registers itself
#   in the Windows Running Object Table while alive), instead of spawning a
#   fresh process. That's what produced the intermittent lag/skips.
#   Strategy:
#     • DispatchEx() — bypasses the ROT, always creates a brand-new process
#     • ALL cleanup moved into finally — runs even if SaveAs/Open throws
#     • After Quit(), actively poll until the WINWORD.EXE PID has fully
#       exited (bounded wait) before returning, instead of assuming it's
#       gone the instant Quit() returns
def _wait_for_process_exit(pid, timeout=8):
    """Block until the given PID has fully exited, or timeout (seconds)."""
    if not pid:
        return
    import win32api, win32con, win32process, time
    try:
        handle = win32api.OpenProcess(win32con.PROCESS_QUERY_INFORMATION, False, pid)
    except Exception:
        return  # already gone
    start = time.time()
    try:
        while time.time() - start < timeout:
            if win32process.GetExitCodeProcess(handle) != win32con.STILL_ACTIVE:
                return
            time.sleep(0.15)
    finally:
        win32api.CloseHandle(handle)


def to_pdf(docx_path, pdf_path):
    word, doc, pid = None, None, None
    try:
        import win32com.client
        import win32process
        import pythoncom
        import gc

        pythoncom.CoInitialize()
        # DispatchEx forces a NEW Word process — never attaches to a
        # previous instance that may still be mid-Quit()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        try:
            pid = win32process.GetWindowThreadProcessId(word.Hwnd)[1]
        except Exception:
            pid = None  # PID capture is best-effort; cleanup still proceeds

        doc = word.Documents.Open(os.path.abspath(docx_path), ReadOnly=True)
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        return True

    except ImportError:
        print("    [!] pywin32 missing — saving .docx instead of PDF.")
        shutil.copy(docx_path, pdf_path.replace(".pdf", ".docx"))
        return False
    except Exception as e:
        print(f"    [!] PDF error: {e}")
        return False
    finally:
        # Cleanup ALWAYS runs now, success or failure
        try:
            if doc is not None:
                doc.Close(SaveChanges=False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        try:
            del doc, word
        except Exception:
            pass
        try:
            import gc
            gc.collect()
        except Exception:
            pass
        _wait_for_process_exit(pid)          # block until the PID is actually gone
        try:
            import pythoncom
            pythoncom.CoUninitialize()
        except Exception:
            pass




# ─── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    print("=" * 62)
    print("  XDrives LP Reporting Automation  —  v4")
    print(f"  {datetime.now():%Y-%m-%d %H:%M:%S}")
    print("=" * 62)

    for p, lbl in [(CAS_FILE, "CAS Excel"), (WORD_TEMPLATE, "Word template")]:
        if not os.path.exists(p):
            print(f"\n[ERROR] {lbl} not found: {p}"); sys.exit(1)

    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    tmp = os.path.join(OUTPUT_FOLDER, "_tmp")
    os.makedirs(tmp, exist_ok=True)

    print(f"\n[1/3] Reading {CAS_FILE} …")
    lps = read_cas(CAS_FILE, CAS_SHEET)
    print(f"      {len(lps)} LP(s) found.")

    pdf_map = {}             # ukey (Number) -> output path, for the final summary
    entity_folders = set()   # tracks distinct entity subfolders created
    used_filenames = {}      # (folder, base_filename) -> count, for collision-safe naming

    print(f"\n[2/3] Generating PDFs → {OUTPUT_FOLDER} (grouped by Entity)")
    for lp in lps:
        lp_number = lp.get("Number", "")
        name      = lp.get("Investor", "Unknown")
        entity    = lp.get("Entity", "Unknown Entity")
        print(f"\n  ▶  [{lp_number}] {name}  —  {entity}")

        # Build token replacements
        tokens = {}
        for tok, col in COLUMN_MAP.items():
            raw = lp.get(col, "")

            if tok == "{{INVESTED_CAPITAL}}":
                tokens[tok] = fmt_capital(raw)
            elif tok == "{{DATE}}":
                tokens[tok] = fmt_date(raw)
            else:
                tokens[tok] = fmt(raw)

        doc = fill_doc(WORD_TEMPLATE, tokens)

        # Per-entity output subfolder, e.g. Output_PDFs/Ram Group Partners, LP/
        entity_folder = os.path.join(OUTPUT_FOLDER, safe_folder_name(entity))
        os.makedirs(entity_folder, exist_ok=True)
        entity_folders.add(entity_folder)

        ukey = unique_key(lp_number)          # Number column — temp .docx name / pdf_map key only
        docx = os.path.join(tmp, f"{ukey}_{QUARTER_LABEL}.docx")

        # "[Entity] - MM.DD.YYYY - [Investor].pdf" — Number is unique, but it
        # isn't part of this filename, so two rows can still share the same
        # entity + date + investor name. Guard against that collision by
        # appending " (2)", " (3)" etc. instead of silently overwriting.
        raw_date     = lp.get("Date", "")
        base_name    = build_pdf_filename(entity, raw_date, name)
        collision_id = (entity_folder, base_name)
        used_filenames[collision_id] = used_filenames.get(collision_id, 0) + 1
        occurrence = used_filenames[collision_id]
        if occurrence > 1:
            stem = base_name[:-4]  # strip ".pdf"
            pdf_filename = f"{stem} ({occurrence}).pdf"
            print(f"     [i] Duplicate name detected for this entity/date — "
                  f"saved as \"{pdf_filename}\" to avoid overwrite.")
        else:
            pdf_filename = base_name
        pdf = os.path.join(entity_folder, pdf_filename)

        doc.save(docx)

        ok = to_pdf(docx, pdf)
        print(f"     {'✅ PDF' if ok else '⚠ DOCX'} → "
              f"{pdf if ok else pdf.replace('.pdf','.docx')}")
        pdf_map[ukey] = pdf if ok else pdf.replace(".pdf", ".docx")

    shutil.rmtree(tmp, ignore_errors=True)

    print(f"\n[3/3] Complete.")
    print(f"      PDFs : {len(pdf_map)}   Entity folders : {len(entity_folders)}")
    print(f"      Output folder : {os.path.abspath(OUTPUT_FOLDER)}")
    print("\n  ✅  Done — PDFs are sorted into their entity subfolders.\n")


if __name__ == "__main__":
    main()