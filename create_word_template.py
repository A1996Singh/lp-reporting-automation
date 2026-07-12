"""
create_word_template.py
========================
Generates LP_Report_Template_v3.docx — the Word template used by
lp_report_automation.py for LP quarterly reporting.

Run this script ONCE to (re)create the template.

LOGO SETUP (do this once before running):
  1. Place your company logo image in the SAME folder as this script.
  2. Set LOGO_PATH below to the filename, e.g.  LOGO_PATH = "logo.png"
  3. Supported formats: PNG, JPG, BMP
  4. Leave LOGO_PATH = "" to skip the logo.

FOOTER TEXT:
  Edit FOOTER_TEXT below to change the contact line in the footer.

Requirements (install once):
  pip install python-docx

Run:
  python create_word_template.py

Output:
  LP_Report_Template_v3.docx  (created in the same folder as this script)
"""

from email.header import Header
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── USER SETTINGS ─────────────────────────────────────────────────────────────
FOOTER_TEXT = (
    "If you have any questions with your statement, "
    "please contact azad.chaudhary@stardom@gmail.com"
)
LOGO_PATH   = "Y_Axis Logo.png"   # filename of your logo file; set "" to skip
LOGO_HEIGHT = Inches(0.5)  # logo display height (width scales automatically)
# ──────────────────────────────────────────────────────────────────────────────

# ── Twip constants (1 inch = 1440 twips) ──────────────────────────────────────
CD = int(3.0   * 1440)   # 4320  — description column
CN = int(1.167 * 1440)   # 1680  — each numeric column (x3 = 5040)
TW = CD + CN * 3         # 9360  — total table width (6.5")

# ── Border specs ──────────────────────────────────────────────────────────────
NONE_B  = None
THIN_B  = {"val": "single", "sz": "4",  "color": "888888"}
MED_B   = {"val": "single", "sz": "8",  "color": "000000"}
THICK_B = {"val": "single", "sz": "14", "color": "000000"}

# ── XML helpers ───────────────────────────────────────────────────────────────

def add_para_bottom_border(para, sz="12", color="000000"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    sz)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def set_para_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"),   str(before))
    sp.set(qn("w:after"),    str(after))
    sp.set(qn("w:line"),     "240")
    sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)

def set_row_height(row, twips, rule="atLeast"):
    trPr = row._tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:trHeight")):
        trPr.remove(old)
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(twips))
    trH.set(qn("w:hRule"), rule)
    trPr.append(trH)

def no_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblB = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:color"), "FFFFFF")
        tblB.append(el)
    tblPr.append(tblB)

def set_table_width(table, twips):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    tblPr.append(lay)

def make_tc(text, width, size_pt=11, bold=False,
            align="left",
            top=None, bottom=None, left=None, right=None,
            sp_before=50, sp_after=50):
    tc   = OxmlElement("w:tc")
    tcPr = OxmlElement("w:tcPr")
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(width))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)
    tcBorders = OxmlElement("w:tcBorders")
    for side, spec in [("top", top), ("bottom", bottom),
                       ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        if spec is None:
            el.set(qn("w:val"),   "none")
            el.set(qn("w:sz"),    "0")
            el.set(qn("w:color"), "FFFFFF")
        else:
            el.set(qn("w:val"),   spec["val"])
            el.set(qn("w:sz"),    spec["sz"])
            el.set(qn("w:color"), spec["color"])
        tcBorders.append(el)
    tcPr.append(tcBorders)
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", 50), ("bottom", 50), ("left", 100), ("right", 100)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)
    vA = OxmlElement("w:vAlign")
    vA.set(qn("w:val"), "center")
    tcPr.append(vA)
    tc.append(tcPr)
    p   = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc  = OxmlElement("w:jc")
    jc.set(qn("w:val"), align)
    pPr.append(jc)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"),   str(sp_before))
    sp.set(qn("w:after"),    str(sp_after))
    sp.set(qn("w:line"),     "240")
    sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)
    p.append(pPr)
    if text:
        r   = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"),    "Calibri")
        rFonts.set(qn("w:hAnsi"),    "Calibri")
        rFonts.set(qn("w:eastAsia"), "Calibri")
        rPr.append(rFonts)
        half = str(int(size_pt * 2))
        for tag in ("w:sz", "w:szCs"):
            el = OxmlElement(tag)
            el.set(qn("w:val"), half)
            rPr.append(el)
        if bold:
            rPr.append(OxmlElement("w:b"))
            rPr.append(OxmlElement("w:bCs"))
        col_el = OxmlElement("w:color")
        col_el.set(qn("w:val"), "000000")
        rPr.append(col_el)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        if text != text.strip():
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        p.append(r)
    tc.append(p)
    return tc


def add_row(table, cells):
    tr = OxmlElement("w:tr")
    for tc in cells:
        tr.append(tc)
    table._tbl.append(tr)
    from docx.table import _Row
    return _Row(tr, table)


# =============================================================================
def build_template():
    doc = Document()

    for sec in doc.sections:
        sec.page_width    = Inches(8.5)
        sec.page_height   = Inches(11)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)
        sec.top_margin    = Inches(0.8)
        sec.bottom_margin = Inches(0.8)

    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # 1. Large title
    tp = doc.add_paragraph()
    tr = tp.add_run("{{ENTITY}}")
    tr.font.name = "Calibri"
    tr.font.size = Pt(28)
    tr.font.color.rgb = RGBColor(0, 0, 0)
    set_para_spacing(tp, 0, 0)
    add_para_bottom_border(tp, sz="12")

    # 2. Three blank gap rows
    for _ in range(3):
        g = doc.add_paragraph()
        set_para_spacing(g, 0, 0)

    # 3. Sub-header (entity name, bold, smaller font)
    sp2 = doc.add_paragraph()
    sr  = sp2.add_run("{{ENTITY}}")
    sr.font.name  = "Calibri"
    sr.font.size  = Pt(11)
    sr.font.bold  = True
    sr.font.color.rgb = RGBColor(0, 0, 0)
    set_para_spacing(sp2, 0, 0)

    # 4. CAS line
    cp = doc.add_paragraph()
    cr = cp.add_run("Capital Account Statement as of {{DATE}}")
    cr.font.name = "Calibri"
    cr.font.size = Pt(11)
    cr.font.color.rgb = RGBColor(0, 0, 0)
    set_para_spacing(cp, 0, 200)

    # 5. Investor info (borderless 2-col table)
    W1   = int(1.6 * 1440)
    W2   = TW - W1
    info = doc.add_table(rows=0, cols=1)
    no_table_borders(info)
    set_table_width(info, TW)
    for lbl, tok in [("Investor:",         "{{INVESTOR_NAME}}"),
                     ("Invested Capital:", "{{INVESTED_CAPITAL}}")]:
        add_row(info, [
            make_tc(lbl, W1, align="left", sp_before=30, sp_after=30),
            make_tc(tok, W2, align="left", sp_before=30, sp_after=30),
        ])

    pre = doc.add_paragraph()
    set_para_spacing(pre, 200, 80)

    # 6. NAV table
    # IMPORTANT: cols=1 is required — add_row() builds cells via raw XML.
    # cols=4 would pre-create phantom cells and break row.cells indexing.
    nav = doc.add_table(rows=0, cols=1)
    no_table_borders(nav)
    set_table_width(nav, TW)

    def dr(desc, qtd, ytd, itd, bold=False, is_end=False, indent=False):
        bot = THICK_B if is_end else THIN_B
        d   = "    " + desc if indent else desc
        row = add_row(nav, [
            make_tc(d,   CD, align="left",  bold=bold, top=THIN_B, bottom=bot, left=THIN_B, right=NONE_B),
            make_tc(qtd, CN, align="right", bold=bold, top=THIN_B, bottom=bot, left=THIN_B, right=NONE_B),
            make_tc(ytd, CN, align="right", bold=bold, top=THIN_B, bottom=bot, left=THIN_B, right=NONE_B),
            make_tc(itd, CN, align="right", bold=bold, top=THIN_B, bottom=bot, left=THIN_B, right=THIN_B),
        ])
        for cell_idx in [1, 2, 3]:
            cell = row.cells[cell_idx]
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name      = "Avenir Next LT Pro Light"
                    r.font.bold      = True
                    r.font.underline = False


    def sl(label):
        add_row(nav, [
            make_tc(label, CD, align="left", top=NONE_B, bottom=NONE_B, left=THIN_B, right=NONE_B, sp_before=120, sp_after=20),
            make_tc("",    CN, align="left", top=NONE_B, bottom=NONE_B, left=THIN_B, right=NONE_B, sp_before=120, sp_after=20),
            make_tc("",    CN, align="left", top=NONE_B, bottom=NONE_B, left=THIN_B, right=NONE_B, sp_before=120, sp_after=20),
            make_tc("",    CN, align="left", top=NONE_B, bottom=NONE_B, left=THIN_B, right=THIN_B, sp_before=120, sp_after=20),
        ])

    def sp():
        r = add_row(nav, [
            make_tc("", CD, top=NONE_B, bottom=NONE_B, left=THIN_B, right=NONE_B, sp_before=0, sp_after=0),
            make_tc("", CN, top=NONE_B, bottom=NONE_B, left=THIN_B, right=NONE_B, sp_before=0, sp_after=0),
            make_tc("", CN, top=NONE_B, bottom=NONE_B, left=THIN_B, right=NONE_B, sp_before=0, sp_after=0),
            make_tc("", CN, top=NONE_B, bottom=NONE_B, left=THIN_B, right=THIN_B, sp_before=0, sp_after=0),
        ])
        set_row_height(r, 220)

    # Header row
    add_row(nav, [
        make_tc("",                CD, align="left",   top=NONE_B, bottom=NONE_B, left=NONE_B, right=NONE_B, sp_before=40, sp_after=40),
        make_tc("QTD",             CN, align="center", top=NONE_B, bottom=MED_B,  left=NONE_B, right=NONE_B, sp_before=40, sp_after=40),
        make_tc("YTD",             CN, align="center", top=NONE_B, bottom=MED_B,  left=NONE_B, right=NONE_B, sp_before=40, sp_after=40),
        make_tc("Since Inception", CN, align="center", top=NONE_B, bottom=MED_B,  left=NONE_B, right=NONE_B, sp_before=40, sp_after=40),
    ])

    # Data rows
    dr("Beginning Net Asset Value",
       "{{QTD_BEGINNING_NAV}}", "{{YTD_BEGINNING_NAV}}", "{{ITD_BEGINNING_NAV}}")
    sp()

    sl("Contributions")
    dr("Contributions - Class A", "-", "-", "{{ITD_CONTRIBUTIONS}}", indent=True)
    sp()

    dr("Change in Unrealized Appreciation / (Depreciation) on Investment",
       "{{QTD_CHANGE_UNREALIZED}}", "{{YTD_CHANGE_UNREALIZED}}", "{{ITD_CHANGE_UNREALIZED}}")
    sp()

    sl("Distributions")
    dr("Return of Capital",
       "{{QTD_RETURN_OF_CAPITAL}}", "{{YTD_RETURN_OF_CAPITAL}}", "{{ITD_RETURN_OF_CAPITAL}}",
       indent=True)
    dr("Return on Capital",
       "{{QTD_RETURN_ON_CAPITAL}}", "{{YTD_RETURN_ON_CAPITAL}}", "{{ITD_RETURN_ON_CAPITAL}}",
       indent=True)
    sp()

    dr("Ending Net Asset Value",
       "{{QTD_ENDING_NAV}}", "{{YTD_ENDING_NAV}}", "{{ITD_ENDING_NAV}}",
       is_end=True)

    # 7. FOOTER
    Footer = doc.sections[0].footer
    for fp in Footer.paragraphs:
        for run in fp.runs:
            run.text = ""

    fp          = Footer.paragraphs[0]
    logo_abs    = os.path.join(os.path.dirname(os.path.abspath(__file__)), LOGO_PATH)
    logo_exists = LOGO_PATH and os.path.exists(logo_abs)

    if logo_exists:
        fp.alignment   = WD_ALIGN_PARAGRAPH.LEFT
        tab_c          = fp.add_run("\t")
        tab_c.font.size = Pt(8)
        text_run        = fp.add_run(FOOTER_TEXT)
        text_run.font.name      = "Calibri"
        text_run.font.size      = Pt(8)
        text_run.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
        tab_r           = fp.add_run("\t")
        tab_r.font.size = Pt(8)
        logo_run        = fp.add_run()
        logo_run.add_picture(logo_abs, height=LOGO_HEIGHT)
        print(f"   Logo added from: {logo_abs}")
    else:
        fp.alignment            = WD_ALIGN_PARAGRAPH.CENTER
        text_run                = fp.add_run(FOOTER_TEXT)
        text_run.font.name      = "Calibri"
        text_run.font.size      = Pt(8)
        text_run.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
        if LOGO_PATH:
            print(f"   Logo file not found: {logo_abs}")
            print(f"   Place '{LOGO_PATH}' in the same folder and rerun.")

    doc.save("LP_Report_Template_v3.docx")
    print("LP_Report_Template_v3.docx created successfully.")


if __name__ == "__main__":
    build_template()
