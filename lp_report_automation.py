"""
XDrives LP Reporting Automation  — v3
=======================================
Excel CAS  →  Word template (per LP)  →  PDF  →  Outlook Draft

Requirements (install once):
    pip install openpyxl python-docx pywin32

Run:
    python lp_report_automation.py

CHANGES FROM v2:
  Line 109  fmt()         — rounds to integer, zero → "-", no .00
  Line 122  fmt_capital() — invested capital as rounded integer currency
  Line 130  fmt_date()    — strips time component, keeps date only
  Line 155  read_cas()    — returns LP number as part of record
  Line 170  safe_name()   — uses LP Number as unique key prefix to avoid
                            duplicate-name DOCX overwrite (Problem 2)
  Line 195  fill_doc()    — unchanged
  Line 218  to_pdf()      — Word COM fully quit+del between each LP so
                            documents are released before next open (Problem 1)
  Line 248  main()        — tokens block uses fmt_date() for {{DATE}} (Problem 3)
"""

import os, re, sys, shutil, openpyxl
from docx import Document
from datetime import datetime

# ─── CONFIGURATION ────────────────────────────────────────────────────────────
CAS_FILE      = r"CAS_Data\test_NAV.xlsx"
CAS_SHEET     = "Summary"
WORD_TEMPLATE = r"LP_Report_Template_v3.docx"
OUTPUT_FOLDER = r"Output_PDFs"
QUARTER_LABEL = "Q1 2026"

DRAFT_SUBJECT = f"Capital Account Statement — {QUARTER_LABEL}"
DRAFT_BODY    = (
    "Dear {{INVESTOR_NAME}},\n\n"
    f"Please find attached your Capital Account Statement for {QUARTER_LABEL}.\n\n"
    "Should you have any questions, please do not hesitate to contact us.\n\n"
    "Best regards,\n"
    "Fund Accounting Team"
)

# ─── COLUMN → TEMPLATE TOKEN MAP ──────────────────────────────────────────────
COLUMN_MAP = {
    "{{INVESTOR_NAME}}":         "Investor",
    "{{ENTITY}}":                "Entity",
    "{{INVESTOR_EMAIL}}":        "Investor Email",
    "{{INVESTOR_CC_EMAIL}}":     "Investor CC Email",
    "{{DATE}}":                  "Date",
    "{{INVESTED_CAPITAL}}":      "Invested Capital",
    "{{ITD_BEGINNING_NAV}}":     "ITD Beginning NAV",
    "{{ITD_CONTRIBUTIONS}}":     "ITD Contributions",
    "{{ITD_CHANGE_UNREALIZED}}": "ITD Change in Unrealized App Dep",
    "{{ITD_RETURN_OF_CAPITAL}}": "ITD Return of Capital",
    "{{ITD_RETURN_ON_CAPITAL}}": "ITD Return on Capital",
    "{{ITD_ENDING_NAV}}":        "ITD Ending NAV",
    "{{YTD_BEGINNING_NAV}}":     "YTD beginning NAV",
    "{{YTD_CONTRIBUTIONS}}":     "YTD Contributions",
    "{{YTD_CHANGE_UNREALIZED}}": "YTD Change in Unrealized App Dep",
    "{{YTD_RETURN_OF_CAPITAL}}": "YTD Return of Capital",
    "{{YTD_RETURN_ON_CAPITAL}}": "YTD Return on Capital",
    "{{YTD_ENDING_NAV}}":        "YTD Ending NAV",
    "{{QTD_BEGINNING_NAV}}":     "QTD Beginning NAV",
    "{{QTD_CONTRIBUTIONS}}":     "QTD Contributions",
    "{{QTD_CHANGE_UNREALIZED}}": "QTD Change in Unrealized App Dep",
    "{{QTD_RETURN_OF_CAPITAL}}": "QTD Return of Capital",
    "{{QTD_RETURN_ON_CAPITAL}}": "QTD Return on Capital ",  # trailing space in Excel
    "{{QTD_ENDING_NAV}}":        "QTD Ending NAV",
}

# ─── HELPERS ──────────────────────────────────────────────────────────────────

# LINE 109 — FIX Problem 3a, 3b & 3d:
#   • Rounds all values to integer (no .00)
#   • Zero / blank / dash / None  →  "-"  (no "$0" shown)
#   • Negative numbers            →  "($2,313)"  not  "-$2,313"
def fmt(v):
    """
    Format a NAV/financial value for display:
      - None / "" / "-" / 0 / 0.0  →  "-"
      - Positive number             →  "$12,345"
      - Negative number             →  "($2,313)"   ← accounting parentheses format
      - Non-numeric string          →  returned as-is
    """
    if v is None or str(v).strip() in ("", "-", "—"):
        return "-"
    try:
        num = float(v)
        # LINE 120 — zero activity shows as "-" not "$0"
        if num == 0:
            return "-"
        # LINE 122 — FIX Problem 3d: negative → accounting parentheses format
        if num < 0:
            return f"(${abs(int(round(num))):,})"
        return f"${int(round(num)):,}"
    except (ValueError, TypeError):
        return str(v)


# LINE 126 — FIX Problem 3b & 3d: Invested Capital rounded integer, negatives in parens
def fmt_capital(v):
    """Invested Capital: rounded integer, zero → '-', negative → '($n)'."""
    if v is None or str(v).strip() in ("", "-", "—"):
        return "-"
    try:
        num = float(v)
        if num == 0:
            return "-"
        if num < 0:
            return f"(${abs(int(round(num))):,})"
        return f"${int(round(num)):,}"
    except (ValueError, TypeError):
        return str(v) if v else "-"


# LINE 130 — FIX Problem 3c: date only, strip any time component
def fmt_date(v):
    """
    Return only the date portion of the CAS date field.
    Handles: datetime objects, 'March 31,2026', '2026-03-31', '2026-03-31 00:00:00'
    """
    if v is None or str(v).strip() == "":
        return ""
    # Already a datetime/date object
    if hasattr(v, "strftime"):
        return v.strftime("%B %d, %Y")          # e.g. "March 31, 2026"
    s = str(v).strip()
    # Try parsing common formats and normalise
    for fmt_str in ("%B %d,%Y", "%B %d, %Y", "%Y-%m-%d %H:%M:%S",
                    "%Y-%m-%d", "%m/%d/%Y", "%d-%b-%Y"):
        try:
            return datetime.strptime(s, fmt_str).strftime("%B %d, %Y")
        except ValueError:
            continue
    # Fallback: strip anything after a space that looks like a time (HH:MM)
    s = re.sub(r"\s+\d{1,2}:\d{2}(:\d{2})?(\.\d+)?$", "", s).strip()
    return s


def read_cas(path, sheet):
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb[sheet]
    rows = list(ws.iter_rows(values_only=True))
    headers = [str(h).strip() if h else "" for h in rows[0]]
    return [{headers[i]: rows[r][i] for i in range(len(headers))}
            for r in range(1, len(rows)) if any(rows[r])]


# LINE 170 — FIX Problem 2: prefix filename with LP Number (col A) so
#            duplicate investor names never produce the same filename.
def safe_name(lp_number, investor_name):
    """
    Build a unique filename-safe string using LP Number + Investor Name.
    e.g.  Number=1, Name="ABC Holdings, LLC"  →  "001_ABC_Holdings_LLC"
    This prevents two LPs with the same investor name from colliding.
    """
    num_str  = str(int(lp_number)).zfill(3) if lp_number else "000"
    name_str = re.sub(r'[\\/*?:"<>|,.]', "_", str(investor_name)).strip()
    return f"{num_str}_{name_str}"


def replace_in_para(para, tokens):
    """Word may split a merge field across runs — rebuild, replace, rewrite."""
    full = "".join(r.text for r in para.runs)
    if any(tok in full for tok in tokens):
        for tok, val in tokens.items():
            full = full.replace(tok, val)
        if para.runs:
            para.runs[0].text = full
            for r in para.runs[1:]:
                r.text = ""


def fill_doc(template_path, tokens):
    doc = Document(template_path)
    for para in doc.paragraphs:
        replace_in_para(para, tokens)
    for tbl in doc.tables:
        for row in tbl.rows:
            for c in row.cells:
                for para in c.paragraphs:
                    replace_in_para(para, tokens)
    for sec in doc.sections:
        for hdr_ftr in (sec.header, sec.footer):
            for para in hdr_ftr.paragraphs:
                replace_in_para(para, tokens)
    return doc


# LINE 218 — FIX Problem 1: Word COM fully released after EACH document.
#   Strategy:
#     • Open a fresh Word.Application per LP
#     • SaveAs PDF
#     • Close document with SaveChanges=False
#     • Quit Word
#     • del the COM object + force garbage-collect so the process exits
#       before the next LP tries to open Word again
def to_pdf(docx_path, pdf_path):
    try:
        import win32com.client
        import pythoncom
        import gc

        pythoncom.CoInitialize()                          # ensure COM initialised on this thread
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False

        doc = word.Documents.Open(
            os.path.abspath(docx_path),
            ReadOnly=True                                 # open read-only to avoid lock
        )
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        doc.Close(SaveChanges=False)                      # close WITHOUT saving changes
        word.Quit()
        del doc, word                                     # release COM references
        gc.collect()                                      # force GC so process exits cleanly
        pythoncom.CoUninitialize()
        return True

    except ImportError:
        print("    [!] pywin32 missing — saving .docx instead of PDF.")
        shutil.copy(docx_path, pdf_path.replace(".pdf", ".docx"))
        return False
    except Exception as e:
        print(f"    [!] PDF error: {e}")
        # Attempt cleanup even on failure
        try:
            word.Quit()
        except Exception:
            pass
        return False


def make_draft(to, cc, subject, body, attachment):
    try:
        import win32com.client
        ol   = win32com.client.Dispatch("Outlook.Application")
        mail = ol.CreateItem(0)
        mail.To = to
        if cc and cc != to:
            mail.CC = cc
        mail.Subject = subject
        mail.Body    = body
        if os.path.exists(attachment):
            mail.Attachments.Add(os.path.abspath(attachment))
        mail.Save()
        return True
    except ImportError:
        print("    [!] pywin32 missing — Outlook draft skipped.")
        return False
    except Exception as e:
        print(f"    [!] Outlook error: {e}")
        return False


# ─── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    print("=" * 62)
    print("  XDrives LP Reporting Automation  —  v3")
    print(f"  {datetime.now():%Y-%m-%d %H:%M:%S}")
    print("=" * 62)

    for p, lbl in [(CAS_FILE, "CAS Excel"), (WORD_TEMPLATE, "Word template")]:
        if not os.path.exists(p):
            print(f"\n[ERROR] {lbl} not found: {p}"); sys.exit(1)

    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    tmp = os.path.join(OUTPUT_FOLDER, "_tmp")
    os.makedirs(tmp, exist_ok=True)

    print(f"\n[1/4] Reading {CAS_FILE} …")
    lps = read_cas(CAS_FILE, CAS_SHEET)
    print(f"      {len(lps)} LP(s) found.")

    pdf_map = {}

    print(f"\n[2/4] Generating PDFs → {OUTPUT_FOLDER}")
    for lp in lps:
        # LINE 248 — use Number column as unique key
        lp_number = lp.get("Number", "")
        name      = lp.get("Investor", "Unknown")
        email     = lp.get("Investor Email", "")
        cc        = lp.get("Investor CC Email", "")
        print(f"\n  ▶  [{lp_number}] {name}")

        # Build token replacements
        tokens = {}
        for tok, col in COLUMN_MAP.items():
            raw = lp.get(col, "")

            if tok == "{{INVESTED_CAPITAL}}":
                # LINE 257 — FIX Problem 3b: rounded integer capital
                tokens[tok] = fmt_capital(raw)

            elif tok == "{{DATE}}":
                # LINE 260 — FIX Problem 3c: date only, no time
                tokens[tok] = fmt_date(raw)

            else:
                # LINE 263 — FIX Problem 3a & 3b: rounded integer, zero → "-"
                tokens[tok] = fmt(raw)

        doc = fill_doc(WORD_TEMPLATE, tokens)

        # LINE 170 FIX: unique filename = LPNumber_InvestorName_Quarter
        sn   = safe_name(lp_number, name)
        docx = os.path.join(tmp,           f"{sn}_{QUARTER_LABEL}.docx")
        pdf  = os.path.join(OUTPUT_FOLDER, f"{sn}_{QUARTER_LABEL}.pdf")

        doc.save(docx)

        # LINE 218 FIX: fresh Word COM instance per LP, fully released after
        ok = to_pdf(docx, pdf)
        print(f"     {'✅ PDF' if ok else '⚠ DOCX'} → "
              f"{pdf if ok else pdf.replace('.pdf','.docx')}")
        pdf_map[name] = pdf if ok else pdf.replace(".pdf", ".docx")

    shutil.rmtree(tmp, ignore_errors=True)

    print(f"\n[3/4] Creating Outlook drafts …")
    drafts = 0
    for lp in lps:
        name  = lp.get("Investor", "Unknown")
        email = lp.get("Investor Email", "")
        cc    = lp.get("Investor CC Email", "")
        att   = pdf_map.get(name, "")
        if not email:
            print(f"  [!] No email for {name} — skipped.")
            continue
        body = DRAFT_BODY.replace("{{INVESTOR_NAME}}", name)
        print(f"  ▶  {name}  →  {email}")
        if make_draft(email, cc, DRAFT_SUBJECT, body, att):
            drafts += 1

    print(f"\n[4/4] Complete.")
    print(f"      PDFs : {len(pdf_map)}   Drafts : {drafts}")
    print(f"      Output folder : {os.path.abspath(OUTPUT_FOLDER)}")
    print("\n  ✅  Review Outlook Drafts before sending.\n")


if __name__ == "__main__":
    main()
